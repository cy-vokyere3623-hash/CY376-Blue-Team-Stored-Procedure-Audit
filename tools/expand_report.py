from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path

path = Path(r"C:\Cybersecurity\Blue team Broni\report\CY376-Blue-Team-Report.docx")
doc = Document(str(path))

extra = [
    (
        "Appendix G — Threat Model Detail",
        [
            "The laboratory threat model assumed an attacker who already had low-privilege database access, for example through a compromised application account. From that foothold, the attacker would search for procedures executable by public, inspect definitions for dynamic SQL, and attempt privilege escalation through EXECUTE AS OWNER. If xp_cmdshell wrappers existed and the feature were later enabled, the attacker could pivot to the operating system.",
            "Blue-team controls interrupt that path at multiple points: remove injectable procedures, eliminate unnecessary elevation, revoke broad execute grants, keep dangerous features disabled, and ensure audit telemetry exists for investigation.",
            "This layered model mirrors defense-in-depth principles taught in network monitoring and auditing courses: prevent where possible, detect what remains, and verify that remediations actually closed the finding.",
        ],
    ),
    (
        "Appendix H — Script Design Rationale",
        [
            "Each audit script was deliberately narrow so beginners could explain it in an interview. Script 01 catalogues assets. Script 02 hunts dangerous code. Script 03 evaluates who can run sensitive objects. Script 04 checks server surface area. Script 05 confirms detective controls.",
            "Separating these concerns makes false confidence less likely. A database can pass a configuration check while still containing injectable procedures. Likewise, clean code with public execute rights remains dangerous.",
            "The remediation script was written to be idempotent where practical: it drops known unsafe objects if present, creates safe replacements, and enables audit objects only when missing. That design supports repeated lab demonstrations without manual cleanup.",
        ],
    ),
    (
        "Appendix I — Interview Defense Notes",
        [
            "If asked why stored procedures are not automatically safe, explain that concatenation into dynamic SQL bypasses parameterization benefits. If asked about EXECUTE AS OWNER, explain that callers inherit owner rights and that least privilege prefers EXECUTE AS CALLER or a dedicated low-privilege account.",
            "If asked how verification was performed, state that the same detection scripts were re-run and returned zero risky rows for code patterns and lab public execute grants, that BlueTeamLabAudit was STARTED, and that remote access showed 0/0.",
            "If asked what would change in production, describe change tickets, peer review, scheduled CIS checks, SIEM forwarding of SQL audit events, and exception management for any feature that must remain enabled.",
        ],
    ),
    (
        "Appendix J — Extended Results Discussion",
        [
            "The critical findings were intentional but representative of real incidents. SQL injection inside procedures often survives because teams assume the procedure boundary is trusted. Privilege elevation through EXECUTE AS OWNER is frequently introduced to make an application work quickly, then forgotten.",
            "The medium finding on missing audit is operationally important. Without telemetry, even perfect prevention eventually fails silently. Enabling SQL Server Audit converts configuration and execution events into evidence that blue teams can triage.",
            "The low finding on remote access illustrates surface-area hygiene. Not every enabled option is an immediate exploit, but unused services and features expand the maintenance and attack surface. Disabling unused features is a standard hardening practice in CIS benchmarks.",
        ],
    ),
    (
        "Appendix K — Mapping Course Learning Outcomes",
        [
            "This project maps to network monitoring, security, and auditing outcomes by requiring students to identify weaknesses, apply standards-based controls, collect evidence, and communicate findings. The printed report trains documentation skill. The GitHub repository trains reproducibility. The presentation and interview train verbal defense of technical decisions.",
            "Completing the closed loop from vulnerable seed to verified remediation demonstrates that security work is not only discovery. It is also change, proof, and communication. Those are the competencies expected of junior blue-team analysts.",
            "Future extensions could include Agent job auditing, linked-server review, login-failure monitoring, and integration with open-source SIEM stacks for dashboarding event 33205 and related SQL audit records.",
        ],
    ),
]


def add_heading(text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)


def add_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(12)


for title, paras in extra:
    add_heading(title, 2)
    for para in paras:
        add_para(para)

doc.save(str(path))
print("Expanded DOCX saved")
