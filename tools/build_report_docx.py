# -*- coding: utf-8 -*-
"""Build CY376 Blue Team Word report for Veronica Okyere."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"C:\Cybersecurity\Blue team Broni")
EVIDENCE = ROOT / "evidence" / "screenshots"
OUT = ROOT / "reports" / "CY376-Blue-Team-Report.docx"

NAME = "Veronica Okyere"
INDEX = "FCM.41.018.206.23"
COURSE = "CY376: Network Monitoring, Security and Auditing"
TOPIC = "Auditing Stored Procedures and Database Objects for Security Weaknesses"
TEAM = "Blue Team"
GITHUB = "https://github.com/cy-vokyere3623-hash/blue-team-stored-procedure-audit"
DATE = "August 2026"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, size=11)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)

    for i in range(1, 4):
        style = styles[f"Heading {i}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        if i == 1:
            style.font.size = Pt(14)
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(10)
        elif i == 2:
            style.font.size = Pt(13)
            style.paragraph_format.space_before = Pt(14)
            style.paragraph_format.space_after = Pt(8)
        else:
            style.font.size = Pt(12)
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(6)


def add_para(doc, text, *, bold=False, italic=False, align="justify", size=12, space_after=8, first_line=True):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space_after)
    if first_line and align == "justify":
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=11, italic=True)
    return p


def add_figure(doc, filename, caption, width=6.0):
    path = EVIDENCE / filename
    if not path.exists():
        add_para(doc, f"[Missing figure file: {filename}]", italic=True, align="center", first_line=False)
        add_caption(doc, caption)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            p = cells[c_i].paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        add_page_number(section)

    configure_styles(doc)

    # ---------------- COVER ----------------
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "KWAME NKRUMAH UNIVERSITY OF SCIENCE AND TECHNOLOGY", align="center", bold=True, size=14, first_line=False, space_after=4)
    add_para(doc, "FACULTY / DEPARTMENT OF COMPUTER SCIENCE", align="center", bold=True, size=12, first_line=False, space_after=18)
    add_para(doc, COURSE, align="center", bold=True, size=13, first_line=False, space_after=18)
    add_para(doc, "END-OF-SEMESTER PROJECT REPORT", align="center", bold=True, size=14, first_line=False, space_after=24)
    add_para(doc, TOPIC, align="center", bold=True, size=14, first_line=False, space_after=12)
    add_para(doc, f"Team Classification: {TEAM}", align="center", size=12, first_line=False, space_after=24)
    add_para(doc, f"Student Name: {NAME}", align="center", size=12, first_line=False, space_after=4)
    add_para(doc, f"Index Number: {INDEX}", align="center", size=12, first_line=False, space_after=4)
    add_para(doc, f"GitHub Repository: {GITHUB}", align="center", size=11, first_line=False, space_after=4)
    add_para(doc, f"Date: {DATE}", align="center", size=12, first_line=False, space_after=4)
    page_break(doc)

    # ---------------- ABSTRACT ----------------
    doc.add_heading("Abstract", level=1)
    add_para(
        doc,
        "Stored procedures and related database objects are common components of enterprise applications, "
        "yet insecure procedure design and weak database configuration remain frequent sources of compromise. "
        "This Blue Team project investigated security weaknesses in stored procedures and database objects "
        "within an isolated Microsoft SQL Server 2022 Express laboratory using the AdventureWorks2022 sample "
        "database. The problem addressed was how defenders can systematically inventory database modules, "
        "detect dangerous coding and permission patterns, assess server surface area, and introduce detective "
        "controls before attackers abuse those weaknesses."
    )
    add_para(
        doc,
        "The methodology combined Centre for Internet Security (CIS) SQL Server guidance, OWASP SQL injection "
        "prevention principles, Microsoft SQL Server Audit capabilities, and NetSPI detective-control practices. "
        "Intentionally vulnerable lab procedures were created to simulate SQL injection through dynamic SQL, "
        "privilege escalation via EXECUTE AS OWNER, an xp_cmdshell wrapper, and excessive public execute grants. "
        "Custom T-SQL audit scripts then inventory modules, hunt dangerous code patterns, review permissions, "
        "check server configuration, and inspect audit status. Six findings were documented, remediated, and "
        "re-verified. After remediation, dangerous code-pattern queries returned zero rows, public execute "
        "grants on lab procedures were removed, SQL Server Audit named BlueTeamLabAudit was started, and "
        "remote access was disabled. The project demonstrates a complete blue-team workflow of inventory, "
        "detection, documentation, remediation, and verification for database object security."
    )
    page_break(doc)

    # ---------------- TOC placeholder ----------------
    doc.add_heading("Table of Contents", level=1)
    add_para(doc, "Update this table in Microsoft Word: right-click and select Update Field.", italic=True, align="left", first_line=False)
    # Insert TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run2 = paragraph.add_run("Right-click here → Update Field to refresh the Table of Contents.")
    set_run_font(run2, size=11, italic=True)
    run._r.append(fldChar3)

    toc_items = [
        "1. Introduction ................................................ 1",
        "2. Literature and Tooling Review ............................... 2",
        "3. Methodology ................................................. 4",
        "4. Implementation .............................................. 6",
        "5. Results and Findings ........................................ 8",
        "6. Analysis and Recommendations ............................... 12",
        "7. Conclusion ................................................. 14",
        "References .................................................... 15",
        "Appendices .................................................... 16",
    ]
    for item in toc_items:
        add_para(doc, item, align="left", first_line=False, space_after=4, size=12)
    page_break(doc)

    # ---------------- 1 INTRODUCTION ----------------
    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Background", level=2)
    add_para(
        doc,
        "Modern applications frequently rely on stored procedures to encapsulate business logic close to data. "
        "From a defensive perspective, this design can improve consistency and reduce direct table access, but "
        "it does not automatically guarantee safety. Stored procedures may still construct dynamic SQL insecurely, "
        "execute under elevated ownership contexts, call high-risk extended procedures, or inherit overly broad "
        "permissions through the public role. When such weaknesses exist, attackers who reach the database layer "
        "can escalate privileges, extract sensitive information, or move toward operating-system command execution."
    )
    add_para(
        doc,
        "Blue team work focuses on reducing exposure and improving detection before incidents occur. In database "
        "environments, that includes inventorying modules, reviewing definitions and permissions, hardening "
        "configuration against CIS benchmarks, and enabling audit trails that record sensitive actions. This "
        "project therefore treats stored procedure and database object auditing as a practical defensive control "
        "activity rather than a purely theoretical exercise."
    )
    doc.add_heading("1.2 Problem Statement", level=2)
    add_para(
        doc,
        "Many student and production labs emphasise application-layer testing while giving less structured "
        "attention to database object hygiene. Without a repeatable audit method, dangerous patterns such as "
        "string-concatenated EXEC statements, EXECUTE AS OWNER, xp_cmdshell wrappers, and public execute grants "
        "can remain unnoticed. The problem addressed by this project is how a blue team analyst can establish a "
        "lab, create representative weaknesses, detect them with evidence, remediate them, and prove the fix."
    )
    doc.add_heading("1.3 Aim and Objectives", level=2)
    add_para(
        doc,
        "The aim of the project was to audit stored procedures and related database objects for security "
        "weaknesses in a controlled SQL Server laboratory and to remediate the identified issues."
    )
    add_para(doc, "Specific objectives were to:", align="left", first_line=False)
    objectives = [
        "Set up SQL Server 2022 Express, SSMS, and AdventureWorks2022 as an isolated lab.",
        "Create intentionally vulnerable stored procedures representing common weaknesses.",
        "Develop and run T-SQL scripts for inventory, code-pattern hunting, permission review, configuration checks, and audit inspection.",
        "Document findings with severity ratings and supporting screenshots.",
        "Remediate the findings and re-verify with the same audit scripts.",
        "Produce a report, evidence package, and GitHub-ready project structure suitable for CY376 submission.",
    ]
    for o in objectives:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(o)
        set_run_font(run, size=12)

    doc.add_heading("1.4 Scope and Limitations", level=2)
    add_para(
        doc,
        "The scope was limited to a local SQL Server Express instance (localhost\\SQLEXPRESS) and the "
        "AdventureWorks2022 database. The project did not attack production systems and did not include "
        "network intrusion detection, SIEM correlation beyond Windows Application log concepts, or full "
        "CIS-CAT automated scoring. Linked-server abuse and SQL Agent job abuse were reviewed conceptually "
        "through literature and NetSPI guidance but were not the primary lab targets. All vulnerable "
        "procedures were lab-only artefacts created for training and then remediated."
    )
    page_break(doc)

    # ---------------- 2 LITERATURE ----------------
    doc.add_heading("2. Literature and Tooling Review", level=1)
    doc.add_heading("2.1 CIS Microsoft SQL Server Benchmarks", level=2)
    add_para(
        doc,
        "The Centre for Internet Security publishes consensus benchmarks for Microsoft SQL Server that define "
        "secure configuration baselines (Center for Internet Security, n.d.). Relevant controls for this project "
        "include surface-area reduction for features such as xp_cmdshell, Ole Automation Procedures, CLR, and "
        "Ad Hoc Distributed Queries; least privilege for roles including public; and auditing and logging "
        "expectations. The benchmark approach shaped the server-configuration script used in Phase 5 of the lab "
        "and provided an industry-recognised reference for classifying configuration findings."
    )
    doc.add_heading("2.2 OWASP SQL Injection Guidance", level=2)
    add_para(
        doc,
        "The OWASP SQL Injection Prevention Cheat Sheet emphasises that stored procedures are not inherently safe "
        "(OWASP Foundation, n.d.). Procedures become unsafe when developers concatenate untrusted input into "
        "dynamic SQL and execute it with EXEC or similar constructs. OWASP therefore recommends parameterized "
        "queries and careful review of sp_executesql, EXEC, and related patterns inside procedure definitions. "
        "This guidance directly informed finding F-001 and the design of script 02-code-patterns.sql."
    )
    doc.add_heading("2.3 Microsoft SQL Server Audit and Catalog Views", level=2)
    add_para(
        doc,
        "Microsoft documentation describes sys.sql_modules as the catalog view that returns definitions for "
        "SQL modules such as procedures, functions, triggers, and views (Microsoft, n.d.-a). Combined with "
        "sys.objects, sys.schemas, and sys.database_permissions, these views enable defensive inventories without "
        "requiring proprietary scanners. Microsoft SQL Server Audit further provides server and database audit "
        "specifications that can log sensitive actions to the Windows Application log (Microsoft, n.d.-b). "
        "Enabling audit for configuration changes and selected procedure execution formed the remediation for "
        "finding F-005."
    )
    doc.add_heading("2.4 NetSPI Detective Controls and Related Practice", level=2)
    add_para(
        doc,
        "NetSPI's SQL Server Detective Control Cheat Sheet documents practical audit specifications for "
        "detecting dangerous activity such as xp_cmdshell execution, OLE automation, external scripts, and "
        "linked-server abuse (Sutherland, n.d.). Although this lab focused on module and permission weaknesses, "
        "NetSPI's work justified creating a named server audit (BlueTeamLabAudit) and associating server and "
        "database specifications rather than relying only on ad hoc queries."
    )
    doc.add_heading("2.5 Tools Used in This Project", level=2)
    add_para(doc, "The following tools supported implementation and evidence collection:", align="left", first_line=False)
    tools = [
        "Microsoft SQL Server 2022 Express Edition as the database engine.",
        "SQL Server Management Studio (SSMS) 21 for connection, script execution, and screenshots.",
        "AdventureWorks2022 sample database as a realistic schema for lab procedures.",
        "Custom T-SQL scripts for inventory, pattern hunting, permissions, configuration, and audit checks.",
        "Windows PowerShell and sqlcmd for automation and verification.",
        "GitHub for repository submission of scripts, evidence, and the final report.",
    ]
    for t in tools:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(t)
        set_run_font(run, size=12)
    page_break(doc)

    # ---------------- 3 METHODOLOGY ----------------
    doc.add_heading("3. Methodology", level=1)
    doc.add_heading("3.1 Laboratory Design", level=2)
    add_para(
        doc,
        "The laboratory was hosted on a Windows workstation named NOVA and used Windows Authentication as "
        "NOVA\\Administrator. SQL Server Express listened as the named instance localhost\\SQLEXPRESS. The "
        "AdventureWorks2022 database provided tables such as Production.Product and HumanResources.Employee for "
        "realistic procedure targets. Figure 1 summarises the topology."
    )
    add_figure(doc, "10-lab-topology-diagram.png", "Figure 1. Blue Team laboratory topology for the SQL Server stored procedure audit.", width=6.2)
    add_figure(doc, "01-ssms-connect-localhost-SQLEXPRESS.png", "Figure 2. SSMS connection dialog for localhost\\SQLEXPRESS using Windows Authentication.", width=5.8)

    doc.add_heading("3.2 Audit Workflow", level=2)
    add_para(
        doc,
        "The workflow followed five defensive stages: (1) prepare the lab and create representative vulnerable "
        "objects; (2) inventory modules; (3) hunt dangerous code and permission patterns; (4) document findings; "
        "and (5) remediate and re-verify. This sequence mirrors standard blue-team practice in which detection "
        "is incomplete unless findings are fixed and tested again."
    )
    add_para(doc, "Table 1 maps project phases to scripts and purpose.", align="left", first_line=False)
    add_table(
        doc,
        ["Phase", "Script / Artefact", "Purpose"],
        [
            ["1 Setup", "setup-vulnerable-procs.sql", "Create intentionally weak lab procedures"],
            ["2 Inventory", "01-inventory.sql", "List modules, EXECUTE AS mode, visibility"],
            ["3 Code hunt", "02-code-patterns.sql", "Find EXEC/xp_cmdshell/EXECUTE AS risks"],
            ["4 Permissions", "03-permissions.sql", "Review public and lab execute grants"],
            ["5 Config", "04-server-config.sql", "CIS surface-area and sysadmin review"],
            ["6 Audit check", "05-audit-check.sql", "List SQL Server Audit status/specs"],
            ["7 Remediate", "remediate-findings.sql", "Fix findings and enable audit"],
            ["8 Verify", "Re-run 02–05", "Confirm risks closed with evidence"],
        ],
    )
    add_caption(doc, "Table 1. Mapping of blue-team phases to laboratory artefacts.")

    doc.add_heading("3.3 Finding Severity Model", level=2)
    add_para(
        doc,
        "Findings were rated Critical, High, Medium, or Low. Critical covered clear SQL injection and OS command "
        "paths. High covered privilege escalation and excessive public grants. Medium covered missing detective "
        "controls. Low covered residual surface-area issues such as remote access when other dangerous features "
        "were already disabled."
    )
    page_break(doc)

    # ---------------- 4 IMPLEMENTATION ----------------
    doc.add_heading("4. Implementation", level=1)
    doc.add_heading("4.1 Creation of Vulnerable Laboratory Procedures", level=2)
    add_para(
        doc,
        "Three intentionally weak procedures were deployed into AdventureWorks2022 to create measurable audit "
        "targets. usp_LabSearchProducts_Unsafe concatenated @ProductName into dynamic SQL and executed it with "
        "EXEC(@sql), illustrating OWASP's warning about unsafe stored procedures. usp_LabGetEmployee_Elevate used "
        "WITH EXECUTE AS OWNER to demonstrate privilege escalation risk. usp_LabRunCommand_Dangerous wrapped "
        "master.dbo.xp_cmdshell. EXECUTE permission on all three was granted to public. Figure 3 shows successful "
        "deployment of the vulnerable objects."
    )
    add_figure(doc, "02-setup-vulnerable-procs-success.png", "Figure 3. Successful creation of intentionally vulnerable lab stored procedures.", width=6.0)

    doc.add_heading("4.2 Inventory and Detection Scripts", level=2)
    add_para(
        doc,
        "Script 01-inventory.sql joined sys.sql_modules, sys.objects, and sys.schemas to list module names, types, "
        "timestamps, EXECUTE AS mode, and whether definitions were visible or encrypted. Script 02-code-patterns.sql "
        "searched module definitions for EXEC(, sp_executesql, xp_cmdshell, OLE, OPENROWSET, and related tokens, "
        "and separately listed procedures with non-null execute_as_principal_id. Script 03-permissions.sql reviewed "
        "public grants and focused EXECUTE grants on usp_Lab% objects. Script 04-server-config.sql queried "
        "sys.configurations for CIS-relevant options and listed sysadmin members. Script 05-audit-check.sql "
        "inspected dm_server_audit_status and audit specification details."
    )
    add_figure(doc, "03-inventory-execute-as-owner-lab-procs.png", "Figure 4. Inventory results highlighting lab procedures and EXECUTE AS OWNER.", width=6.0)

    doc.add_heading("4.3 Remediation Implementation", level=2)
    add_para(
        doc,
        "Remediation was implemented in lab\\remediate-findings.sql. Unsafe and elevated procedures were dropped. "
        "Safer replacements usp_LabSearchProducts_Safe and usp_LabGetEmployee_Safe used static parameterized-style "
        "logic without EXECUTE AS OWNER. The xp_cmdshell wrapper was removed without replacement. Public execute "
        "grants on lab procedures were revoked. Remote access was set to 0. A server audit BlueTeamLabAudit was "
        "created to the Application log with a server audit specification for AUDIT_CHANGE_GROUP and "
        "SERVER_OPERATION_GROUP, plus a database audit specification monitoring execute on the safe lab procedures."
    )
    add_figure(doc, "04-remediate-findings-run.png", "Figure 5. Execution of the remediation script against AdventureWorks2022.", width=6.0)
    page_break(doc)

    # ---------------- 5 RESULTS ----------------
    doc.add_heading("5. Results and Findings", level=1)
    doc.add_heading("5.1 Summary of Findings", level=2)
    add_para(
        doc,
        "The audit produced six findings. Table 2 summarises identifier, severity, status, and object. All findings "
        "were remediated and verified before final reporting."
    )
    add_table(
        doc,
        ["ID", "Severity", "Status", "Object / Area", "Issue"],
        [
            ["F-001", "Critical", "Fixed", "usp_LabSearchProducts_Unsafe", "Dynamic SQL injection pattern"],
            ["F-002", "High", "Fixed", "usp_LabGetEmployee_Elevate", "EXECUTE AS OWNER elevation"],
            ["F-003", "Critical", "Fixed", "usp_LabRunCommand_Dangerous", "xp_cmdshell wrapper"],
            ["F-004", "High", "Fixed", "usp_Lab* / public", "Public EXECUTE grants"],
            ["F-005", "Medium", "Fixed", "SQL Server Audit", "No audit configured initially"],
            ["F-006", "Low", "Fixed", "remote access", "Remote access enabled"],
        ],
    )
    add_caption(doc, "Table 2. Consolidated findings from the stored procedure and database object audit.")
    add_figure(doc, "09-findings-template-documented.png", "Figure 6. Documented findings register with severity, status, and verification notes.", width=5.8)

    doc.add_heading("5.2 Before Remediation Evidence", level=2)
    add_para(
        doc,
        "Before remediation, code-pattern hunting returned usp_LabRunCommand_Dangerous and "
        "usp_LabSearchProducts_Unsafe, while the elevation query returned usp_LabGetEmployee_Elevate with "
        "EXECUTE AS OWNER. Permission review showed public EXECUTE GRANT on all three lab procedures. These "
        "results confirmed that the intentionally weak objects were detectable with the audit scripts."
    )
    add_figure(doc, "11-code-patterns-before-remediation.png", "Figure 7. SSMS results before remediation: dangerous code patterns and EXECUTE AS OWNER.", width=6.0)
    add_figure(doc, "12-permissions-public-execute-before.png", "Figure 8. SSMS results before remediation: public EXECUTE grants on lab procedures.", width=6.0)

    doc.add_heading("5.3 After Remediation Evidence", level=2)
    add_para(
        doc,
        "After remediation, script 02-code-patterns.sql returned zero rows for both dangerous definitions and "
        "elevated execute-as procedures. Server configuration checks showed xp_cmdshell, Ole Automation Procedures, "
        "CLR, Ad Hoc Distributed Queries, and remote access disabled (value and value_in_use = 0), with the sa "
        "account disabled. Script 05-audit-check.sql showed BlueTeamLabAudit in STARTED state with enabled server "
        "and database audit specifications."
    )
    add_figure(doc, "05-code-patterns-after-0-rows.png", "Figure 9. Post-remediation code-pattern hunt returning zero risky rows.", width=6.0)
    add_figure(doc, "07-server-config-cis-surface-area.png", "Figure 10. CIS surface-area configuration results after hardening.", width=6.0)
    add_figure(doc, "08-sql-audit-BlueTeamLabAudit-started.png", "Figure 11. BlueTeamLabAudit started with server and database audit specifications.", width=6.0)
    add_figure(doc, "06-permissions-public-role-audit.png", "Figure 12. Permission audit script execution used during review and verification.", width=6.0)
    page_break(doc)

    # ---------------- 6 ANALYSIS ----------------
    doc.add_heading("6. Analysis and Recommendations", level=1)
    doc.add_heading("6.1 Interpretation of Results", level=2)
    add_para(
        doc,
        "The before-and-after evidence demonstrates that common database weaknesses are both detectable and "
        "fixable when defenders use catalog views and a structured checklist. F-001 and F-003 were critical "
        "because they map to direct data compromise and potential host command execution paths. Even though "
        "xp_cmdshell was disabled at the server level during much of the lab, retaining a wrapper procedure and "
        "public execute rights would become immediately dangerous if the feature were re-enabled. F-002 showed "
        "that ownership chaining through EXECUTE AS OWNER can silently expand caller privilege. F-004 confirmed "
        "that public grants convert object-level bugs into broadly reachable attack surfaces. F-005 showed that "
        "without SQL Server Audit, configuration and procedure abuse may leave insufficient forensic trail. F-006 "
        "was lower severity in this Express lab but still aligned with CIS surface-area reduction."
    )
    doc.add_heading("6.2 Recommendations", level=2)
    add_para(doc, "Based on the findings, the following recommendations are made for similar environments:", align="left", first_line=False)
    recs = [
        "Prohibit dynamic SQL built by string concatenation; require parameterized sp_executesql or static SQL where possible.",
        "Avoid EXECUTE AS OWNER unless a documented least-privilege impersonation account is required and reviewed.",
        "Do not grant EXECUTE on application procedures to public; use dedicated application roles.",
        "Keep xp_cmdshell, Ole Automation, and related high-risk features disabled unless an exception is approved.",
        "Enable SQL Server Audit (or equivalent) for audit-change, server-operation, and sensitive procedure execution events.",
        "Re-run inventory and code-pattern scripts after every significant schema deployment as a regression check.",
        "Store audit scripts and findings under version control so blue-team work remains repeatable and reviewable.",
    ]
    for r in recs:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(r)
        set_run_font(run, size=12)

    doc.add_heading("6.3 Practical Implications for Blue Teams", level=2)
    add_para(
        doc,
        "For blue teams, the main lesson is that database security monitoring should include object definitions "
        "and grants, not only failed logins. Catalog queries are low-cost controls that can be scheduled or run "
        "during change windows. Combining preventive hardening with detective audit specifications creates "
        "defence in depth: unsafe objects are harder to introduce, and attempts to alter audit posture or execute "
        "sensitive procedures are more likely to leave evidence."
    )
    page_break(doc)

    # ---------------- 7 CONCLUSION ----------------
    doc.add_heading("7. Conclusion", level=1)
    add_para(
        doc,
        "This project completed a Blue Team audit of stored procedures and database objects on SQL Server 2022 "
        "Express using AdventureWorks2022. The work established a controlled laboratory, introduced representative "
        "weaknesses, and applied a multi-script audit method grounded in CIS, OWASP, Microsoft, and NetSPI "
        "guidance. Six findings spanning SQL injection patterns, privilege escalation, dangerous feature wrappers, "
        "excessive permissions, missing audit, and remote access were documented with SSMS evidence."
    )
    add_para(
        doc,
        "Remediation removed unsafe objects, introduced safer replacements where appropriate, revoked public "
        "execute grants, disabled remote access, and started BlueTeamLabAudit. Re-verification showed zero "
        "dangerous code-pattern rows and confirmed hardened configuration and active audit specifications. The "
        "outcome is a demonstrable defensive lifecycle: inventory, detect, document, remediate, and verify."
    )
    add_para(
        doc,
        "Future work could extend the same method to linked servers, SQL Agent jobs, encrypted modules, and "
        "automated CIS benchmark scoring, and could forward Windows Application log event 33205 into a SIEM for "
        "continuous monitoring. Within the CY376 scope, however, the project objectives were met and no open "
        "findings remain in the laboratory."
    )
    page_break(doc)

    # ---------------- REFERENCES ----------------
    doc.add_heading("References", level=1)
    refs = [
        "Center for Internet Security. (n.d.). CIS Microsoft SQL Server benchmarks. https://www.cisecurity.org/benchmark/microsoft_sql_server",
        "Microsoft. (n.d.-a). sys.sql_modules (Transact-SQL). Microsoft Learn. https://learn.microsoft.com/en-us/sql/relational-databases/system-catalog-views/sys-sql-modules-transact-sql",
        "Microsoft. (n.d.-b). SQL Server audit action groups and actions. Microsoft Learn. https://learn.microsoft.com/en-us/sql/relational-databases/security/auditing/sql-server-audit-action-groups-and-actions",
        "Microsoft. (n.d.-c). AdventureWorks sample databases. Microsoft Learn. https://learn.microsoft.com/en-us/sql/samples/adventureworks-install-configure",
        "OWASP Foundation. (n.d.). SQL injection prevention cheat sheet. https://cheatsheetseries.owasp.org/cheatsheets/SQL_Injection_Prevention_Cheat_Sheet.html",
        "Sutherland, S. (n.d.). SQL Server detective control cheat sheet. NetSPI PowerUpSQL Wiki. https://github.com/NetSPI/PowerUpSQL/wiki/SQL-Server-Detective-Control-Cheat-Sheet",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-1.25)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(ref)
        set_run_font(run, size=12)
    page_break(doc)

    # ---------------- APPENDICES ----------------
    doc.add_heading("Appendix A. Project Folder Structure", level=1)
    add_para(doc, "The GitHub-oriented project layout used for submission is summarised below.", align="left", first_line=False)
    structure = (
        "Blue team Broni/\n"
        "  README.md\n"
        "  scripts/01-inventory.sql ... 05-audit-check.sql\n"
        "  lab/setup-vulnerable-procs.sql\n"
        "  lab/remediate-findings.sql\n"
        "  docs/findings-template.md\n"
        "  evidence/screenshots/\n"
        "  reports/CY376-Blue-Team-Report.docx\n"
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(structure)
    set_run_font(run, name="Consolas", size=10)

    doc.add_heading("Appendix B. Key Script Excerpts", level=1)
    add_para(
        doc,
        "The following excerpt from the vulnerable search procedure illustrates the unsafe concatenation pattern "
        "identified as F-001:",
        align="left",
        first_line=False,
    )
    code1 = (
        "SET @sql = N'SELECT ProductID, Name, ProductNumber\\n"
        "             FROM Production.Product\\n"
        "             WHERE Name LIKE ''%' + @ProductName + N'%''';\\n"
        "EXEC(@sql);  -- UNSAFE"
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code1.replace("\\n", "\n"))
    set_run_font(run, name="Consolas", size=10)

    add_para(
        doc,
        "The remediation replaced unsafe objects with safer procedures and enabled BlueTeamLabAudit. Full scripts "
        "are included in the repository under scripts\\ and lab\\ rather than duplicated in full here to keep the "
        "main body readable, as required by the submission guidelines for appendices versus core narrative.",
        first_line=True,
    )

    doc.add_heading("Appendix C. Evidence Screenshot Index", level=1)
    add_table(
        doc,
        ["Figure", "File", "Description"],
        [
            ["1", "10-lab-topology-diagram.png", "Lab topology"],
            ["2", "01-ssms-connect-...", "SSMS connection"],
            ["3", "02-setup-vulnerable-...", "Vulnerable procs created"],
            ["4", "03-inventory-...", "Inventory / EXECUTE AS OWNER"],
            ["5", "04-remediate-...", "Remediation run"],
            ["6", "09-findings-template-...", "Findings register"],
            ["7", "11-code-patterns-before-...", "Before: code patterns"],
            ["8", "12-permissions-...-before", "Before: public EXECUTE"],
            ["9", "05-code-patterns-after-...", "After: 0 risky rows"],
            ["10", "07-server-config-...", "CIS surface area"],
            ["11", "08-sql-audit-...", "Audit started"],
            ["12", "06-permissions-...", "Permissions audit view"],
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Exists: {OUT.exists()} Size: {OUT.stat().st_size}")


if __name__ == "__main__":
    build()
